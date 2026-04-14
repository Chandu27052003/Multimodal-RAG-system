import logging
from pathlib import Path

from docx import Document

from .base import BaseParser

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):

    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def parse(self, file_path: Path) -> dict:
        doc = Document(str(file_path))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        tables_text: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                tables_text.append(" | ".join(row_data))

        if tables_text:
            full_text += "\n\n[Tables]\n" + "\n".join(tables_text)

        core = doc.core_properties
        return {
            "text": full_text,
            "pages": None,
            "page_count": None,
            "extra": {
                "author": core.author,
                "created": str(core.created) if core.created else None,
                "modified": str(core.modified) if core.modified else None,
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
            },
        }
